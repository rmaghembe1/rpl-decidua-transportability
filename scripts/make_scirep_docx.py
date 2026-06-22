from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

INPUT = Path("manuscript/manuscript_full_draft_v1_1_scientific_reports_submission_text.md")
OUTPUT = Path("manuscript/docx/RPL_decidua_scientific_reports_submission_v1_1.docx")

def set_cell_text(cell, text):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)

def style_document(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, bold in [
        ("Heading 1", 14, True),
        ("Heading 2", 12, True),
        ("Heading 3", 11, True),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        footer = section.footer
        if footer.paragraphs:
            add_page_number(footer.paragraphs[0])

def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    return p

def add_title_line(doc, text, bold=False, italic=False, size=11, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def parse_table(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        parts = [x.strip() for x in line.strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", x or "") for x in parts):
            continue
        rows.append(parts)
    return rows

def add_markdown_table(doc, table_lines):
    rows = parse_table(table_lines)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        if i == 0:
            set_repeat_table_header(table.rows[i])
        for j in range(ncols):
            text = row[j] if j < len(row) else ""
            set_cell_text(table.cell(i, j), text)
            if i == 0:
                for run in table.cell(i, j).paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()

def is_front_matter_line(line):
    return (
        line.startswith("Short title:")
        or line.startswith("Article type:")
        or line.startswith("Reuben S. Maghembe")
        or line.startswith("1Department")
        or line.startswith("2Department")
        or line.startswith("*Corresponding author:")
        or line.startswith("Email:")
    )

def main():
    if not INPUT.exists():
        raise SystemExit(f"Input file not found: {INPUT}")

    text = INPUT.read_text(encoding="utf-8")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("---\n", "")

    doc = Document()
    style_document(doc)

    lines = text.splitlines()
    i = 0
    title_done = False

    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        if not title_done:
            add_title_line(doc, line.strip(), bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
            title_done = True
            i += 1
            continue

        if is_front_matter_line(line):
            add_title_line(doc, line.strip(), size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
            i += 1
            continue

        if line.startswith("# "):
            heading = line[2:].strip()
            doc.add_heading(heading, level=1)
            i += 1
            continue

        if line.startswith("## "):
            heading = line[3:].strip()
            doc.add_heading(heading, level=2)
            i += 1
            continue

        if line.startswith("### "):
            heading = line[4:].strip()
            doc.add_heading(heading, level=3)
            i += 1
            continue

        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        if re.match(r"^\d+\.\s+", line):
            add_paragraph(doc, line.strip())
            i += 1
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line[2:].strip())
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            i += 1
            continue

        add_paragraph(doc, line.strip())
        i += 1

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")

if __name__ == "__main__":
    main()
